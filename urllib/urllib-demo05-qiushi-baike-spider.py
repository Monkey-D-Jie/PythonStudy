import os
import urllib.request
import  re
from typing import Any, Union, List

from docx import Document
from docx.opc.oxml import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import requests
from bs4 import BeautifulSoup
'''
demo功能简介说明
从 糗事百科 扒取段子，
并存放至本地
'''




def parse_save_data_by_pat(page_num: int):
    pat_parse_data= []
    for i in  range(1,page_num):
        url = f"https://qiushidabaike.com/text_{i}.html"
        # Mine
        # 根据网页返回的内容，设定能过滤内容的正则表达式
        # <div>.*?</div>.*?</dd>
        cur_url_data = urllib.request.urlopen(url).read().decode("utf-8", "ignore")
        pat = '<dd class="content">(.*?)</dd>'
        find_baike_url = re.compile(pat, re.S).findall(cur_url_data)
        clean_content = clean_html_content(find_baike_url)
        pat_parse_data.extend(clean_content)
        print(f"pat------>>>第{i}页的内容已扒取完毕")
    save_path = "H:\\myIdeaWorkSpace\\PythonStudy\\testFile\\baike-content\\pat_qiushi_baike.docx"
    save_to_docx(pat_parse_data, save_path)


def clean_html_content(raw_contents: Union[str, List[str]]) -> List[str]:
    """增强版HTML清理函数，支持处理单个字符串或字符串列表

    Args:
        raw_contents: 可以是单个HTML字符串或包含多个HTML字符串的列表
           示例输入: "<p>内容</p>" 或 ["<div>文本1</div>", "<span>文本2</span>"]

    Returns:
        清理后的纯文本列表，如 ["内容", "文本1 文本2"]
    """
    # 统一输入格式为列表
    contents_list = [raw_contents] if isinstance(raw_contents, str) else raw_contents

    cleaned_results = []

    for html_str in contents_list:
        # 类型安全检查
        if not isinstance(html_str, str):
            cleaned_results.append("")
            continue

        # 核心清理流程
        soup = BeautifulSoup(html_str, 'html.parser')

        # 移除干扰标签
        for tag in soup(['script', 'style', 'meta', 'link', 'noscript']):
            tag.decompose()

        # 转换换行标签
        for br in soup.find_all('br'):
            br.replace_with('\n')

        # 获取结构化文本
        text = soup.get_text(separator='\n', strip=True)

        # 二次清理残留标签
        text = re.sub(r'</?[a-z][^>]*>', '', text)

        # 规范化空白字符（保留换行）
        text = re.sub(r'[ \t\r]+', ' ', text)          # 压缩水平空白
        text = re.sub(r'\n\s*\n', '\n\n', text)         # 保留段落间距
        text = text.translate(str.maketrans({          # 特殊字符替换
            '\xa0': ' ',   # &nbsp;
            '\u200b': ''    # 零宽空格
        }))
        cleaned_results.append(text.strip())

    return cleaned_results

def parse_save_data_by_soup(page_num: int):
    soup_parse_data=[]
    for j in  range(1,page_num):
        url = f"https://qiushidabaike.com/text_{j}.html"
        # 替换原有正则匹配部分
        cur_url_data = urllib.request.urlopen(url).read().decode("utf-8", "ignore")
        soup = BeautifulSoup(cur_url_data, 'lxml')
        #target_dds = soup.find_all('dd', class_='content')
        contents = [dd.get_text('\n', strip=True)
                    for dd in soup.select('dd.content')]  # CSS选择器写法
        # 处理分页场景时添加
        contents = list(filter(None, contents))  # 过滤空内容
        # for i in range(1, len(contents)):
        #     cur_content  = contents[i].replace(">>查看更多", "")
        #     print(str(i) + ":" + cur_content + "\n")
        #     soup_parse_data.append(cur_content)
        soup_parse_data.extend(contents)
        print(f"soup------>>>第{j}页的内容已扒取完毕")
    save_path = "H:\\myIdeaWorkSpace\\PythonStudy\\testFile\\baike-content\\soup_qiushi_baike.docx"
    save_to_docx(soup_parse_data, save_path)

def ensure_dir_exists(file_path: str) -> None:
    """确保文件路径中的目录存在"""
    dir_path = os.path.dirname(file_path)
    if dir_path and not os.path.exists(dir_path):
        os.makedirs(dir_path, exist_ok=True)  # 自动创建多级目录
        print(f"已创建目录: {dir_path}")


def save_to_docx(contents: list, save_path: str) -> None:
    """将内容保存到Word文档

    Args:
        contents: 要保存的内容列表（每个元素为一个段落）
        save_path: 文档保存路径（如：'H:/data/qiushi.docx'）
    """
    try:
        # 保存前，先创建目录
        ensure_dir_exists(save_path)

        doc = Document()

        # 设置默认字体（可选）
        # 关键设置：强制定义中文字体（兼容Windows/Linux）
        doc.styles['Normal'].font.name = '微软雅黑'
        # doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(10.5)

        # 处理HTML实体转义（新增）
        from html import unescape
        cleaned_contents = [unescape(c) for c in contents]

        # 添加内容
        for idx, content in enumerate(cleaned_contents, 1):
            # 添加标题（可选）
            title = doc.add_paragraph()
            title.add_run(f'段子 {idx}').bold = True
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 修改标题样式
            title.style = 'Heading1'

            # 添加页眉页脚
            section = doc.sections[0]
            header = section.header
            header.paragraphs[0].text = "糗事百科精选"

            # 设置页面边距
            section.left_margin = Inches(1.5)  # 需要导入 Inches

            # 添加正文内容并保留换行
            for line in content.split('\n'):
                p = doc.add_paragraph()
                p.add_run(line.strip())

            # 添加段落间距（可选）
            doc.add_paragraph()

        doc.save(save_path)
        print(f"成功保存 {len(contents)} 条内容到 {save_path}")
    except Exception as e:
        print(f"保存文档失败: {str(e)}")


if __name__ == "__main__":
    # 设置统一的请求头
    news_header = ("User-Agent","Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36")
    # 全局添加header
    opener = urllib.request.build_opener()
    opener.addheaders = [news_header]
    # 全局安装opener
    urllib.request.install_opener(opener)
    # parse_save_data_by_pat(2)
    parse_save_data_by_soup(10)
